from selenium import webdriver
import time
import os
import requests
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as Ec
import matplotlib.pyplot as plt
from bs4 import BeautifulSoup
import json
import re
import docx
class superstartload(object):
    def __init__(self,url):
        self.doc=docx.Document()
        self.driver=webdriver.Chrome(r"E:\python脚本\超星挂机\chromedriver.exe")
        self.driver.get(url)
        self.driver.implicitly_wait(5)
        self.item=0
        self.chapter=1
    def loadweb_nocheckcode(self,user,passward):
        userload=self.driver.find_element_by_xpath(r'//input[@id="unameId"]').send_keys(user)
        pwdload=self.driver.find_element_by_xpath(r'//input[@id="passwordId"]').send_keys(passward)
        picweb=self.driver.find_element_by_xpath(r'//img[@id="numVerCode"]')
        self.driver.find_element_by_id("selectSchoolA").click()
        self.driver.find_element_by_id('searchSchool1').send_keys("湖南文理学院")
        self.driver.find_element_by_class_name('zw_t_btn').click()
        time.sleep(0.5)
        self.driver.find_element_by_id('1820').click()
    def savecookies(self):
        cookies=self.driver.get_cookies()   
        jsonCookies = json.dumps(cookies)
        os.chdir("E:\python脚本\超星挂机\savecookies")
        with open('cookiessave.json', 'w') as f:
            f.write(jsonCookies)
    def readcookies(self):
        os.chdir("E:\python脚本\超星挂机\savecookies")
        with open('cookiessave.json', 'r') as f:  
            listCookies=json.loads(f.read())      
            return listCookies
    def savefile(self,name,data):
        os.chdir(r"E:\python脚本\超星挂机\data")
        with open(name,'w') as f:
            f.write(data)
    def readfile(self,name):
        os.chdir(r"E:\python脚本\超星挂机\data")
        with open(name,'r') as f:
            data=f.read()
        return data
    def loadcookies(self,cookies):
        for cookie in cookies:
            self.driver.add_cookie(cookie)
        self.driver.refresh()
    def saveweb(self,course):
        iframe=self.driver.find_element_by_xpath(r'//iframe[@name="frame_content"]')
        self.driver.switch_to.frame(iframe)
        pattern=r"//a[@title="+ r'"' + course + r'"' +']'
        wait=WebDriverWait(self.driver,10)
        try:
            webdriver=wait.until(Ec.presence_of_all_elements_located((By.XPATH,pattern)))
        except:
            print('未找到元素！')
            pass
        html=self.driver.find_element_by_xpath(pattern).get_attribute('href')
        self.savefile('web.txt',html)
    def readweb(self):
        return self.readfile('web.txt')
    def str2times(self,listenstr):
        listenstr+=":"
        tempstr=""
        mode=1
        times=0
        for listentemp in listenstr:
            if listentemp == ':':
                times+=int(tempstr)+int(tempstr)*mode*59
                tempstr=""
                mode=0
            else:
                empstr+=listentemp
        return times
    def __add_picture(self,url,operate):
        headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/74.0.3729.131 Safari/537.36'}
        content=requests.get(url,headers=headers).content
        with open('buf.png','wb') as f:
            f.write(content)
        operate.add_picture('E:\\python脚本\\超星挂机\\data\\'+'buf.png')

    def __add_content(self,content):
        self.doc.add_paragraph(content)

    def __check_question_type(self,content):
        if re.findall('【单选题】',content) != []:
            return 1
        elif re.findall('【多选题】',content) != []:
            return 1
        elif re.findall('【填空题】',content) != []:
            return 2
        elif re.findall('【判断题】',content) != []:
            return 3
        else:
            print('未匹配到选项，参考:'+content)
            return 0
    def __write_title(self,doct):
        self.item+=1
        title=doct.find(attrs={'class':'Zy_TItle clearfix'})
        contents=title.p.text
        operate=self.doc.add_paragraph(str(self.item)).add_run()
        if self.item == 207:
            a=12342
        if title.p.find(name='img') != None:
            contents=str(title.p)
            head=re.compile('p>(.*?)<img')
            tails=re.compile('[(gif)|(pig)]"/>([^<]*)?<')
            headstr=head.search(contents).group(1)
            tailstrs=tails.findall(contents)
            operate.add_text(headstr)
            index=0
            for url in title.p.find_all(name='img'):
                self.__add_picture(url.attrs['src'],operate)
                if tailstrs == []:
                    continue
                else:
                    operate.add_text(tailstrs[index])
                index+=1
        else:
            operate.add_text(contents)
        return self.__check_question_type(title.text)
        

    def __write_choose(self,doct,questiontype):
        if questiontype == 1:
            chooseparent=doct.find(attrs={'class':'Zy_ulTop'})
            choosechild=chooseparent.find_all(attrs={'class':'clearfix'})
            for choose in choosechild:
                operate=self.doc.add_paragraph().add_run()
                option=choose.i.text
                content=choose.a.text
                if choose.a.find(name='img') !=None:
                    contents=str(choose.a)
                    head=re.compile('p>([^<]*)?<img')
                    tails=re.compile('[(gif)|(png)]"/>([^<]*)?</p')
                    if head.search(contents) != None:
                        headstr=head.search(contents).group(1)
                    operate.add_text(option)
                    tailstrs=tails.findall(contents)
                    index=0
                    for url in choose.a.find_all(name='img'):
                        self.__add_picture(url.attrs['src'],operate)
                        if tailstrs !=[]:
                            operate.add_text(tailstrs[index])
                        index+=1
                else  :
                    operate.add_text(option+content)
    def __write_answer(self,doct,questtype):
        if questtype == 2:
            answercontent=doct.find(attrs={'class':'Py_tk'})
            self.doc.add_paragraph(answercontent.text.replace('\n',''))
        else :
            answercontent=doct.find(attrs={'class':'Py_answer clearfix'})
            self.doc.add_paragraph(answercontent.span.text.replace('\n',''))
    def save_doc(self,filename):
        self.doc.save('E:\\python脚本\\超星挂机\\data\\'+filename+'.docx')
    def intomain(self,url,chapter): 
        self.__intomainweb(url,chapter)
    #用途：进入对应章节主页面，并返回根目录,如果根目录不存在，则返回[]
    def __intomainweb(self,url,chapter):
        self.driver.switch_to.default_content()
        self.driver.get(url)
        soup=BeautifulSoup(self.driver.page_source,'lxml')
        temp=soup.find(class_="chapterNumber",text=chapter)
        html=temp.parent.a['href']
        html='https://mooc1-1.chaoxing.com'+html
        self.driver.get(html)
#        print(self.driver.page_source)
        root=list()
        for temp in range(1,10):
            pattern='dct'+str(temp)
                #第一次等待temp载入
            if temp==1:
                wait=WebDriverWait(self.driver,10)
                try:
                    wait.until(Ec.presence_of_element_located((By.ID,pattern)))#as value write error  except catch word error
                except:
                    return []
            try:
                root.append(self.driver.find_element_by_id(pattern))
            except:
                print("当前有"+str(temp-1)+"个子页面")
                break
        return root
    #用途：进入根目录，并返回下面的视屏列表//PPT列表
    def __intosubweb(self,root):
        root.click()
        wait=WebDriverWait(self.driver,10)
        wait.until(Ec.presence_of_element_located((By.ID,"iframe")))
            #document 定位有误：
            #此处有坑，两个iframe
        iframeroot=self.driver.find_element(By.XPATH,r'//iframe[@id="iframe"]')
        time.sleep(1)
        self.driver.switch_to.frame(iframeroot)
        alliframe=self.driver.find_elements_by_tag_name('iframe')
        count=0
        for ses in alliframe:
            count+=1
        print("当前有"+str(count)+"个子任务")
        return alliframe
    def catchcontent(self):
        parentiframe=self.driver.find_element_by_id("iframe")
        self.driver.switch_to.frame(parentiframe)
        childiframe=self.driver.find_element_by_id("ext-gen1039").find_element_by_tag_name("iframe")
        self.driver.switch_to.frame(childiframe)
        iframe=self.driver.find_element_by_id("frame_content")
        self.driver.switch_to.frame(iframe)
        soup=BeautifulSoup(self.driver.page_source,'lxml')
        docts=soup.find_all(attrs={'class':"TiMu"})
        self.doc.add_paragraph('章节'+str(self.chapter))
        for doct in docts:
            qusetion_type=self.__write_title(doct)
            self.__write_choose(doct,qusetion_type)
            self.__write_answer(doct,qusetion_type)
        self.driver.switch_to.default_content()
        self.chapter+=1
indexroot='http://i.mooc.chaoxing.com/space/index?t=1574264821050'
#web.listenclass(studyroot,16)
web=superstartload('http://passport2.chaoxing.com/login?fid=&refer=http://i.mooc.chaoxing.com')
command=input()
todyroot=['1.7','2.10','3.10','4.5','5.7','6.5','7.9','8.3']
if command == 'rush':
    web.loadweb_nocheckcode('201716010116','q1972821')
    input()
    web.savecookies()
    web.saveweb('自动控制原理/MATLAB基础及应用（第4期19下期）')
    studyroot=web.readweb()
    web.intomain(studyroot,todyroot)
    web.catchcontent()
    web.save_doc('pytry')
if command == 'load':
    for root in todyroot:
        web.loadcookies(web.readcookies())
        web.intomain(web.readweb(),root)
        web.catchcontent()
    web.save_doc('superstart')
if command == 'exit':
    print('退出当前程序')
    
